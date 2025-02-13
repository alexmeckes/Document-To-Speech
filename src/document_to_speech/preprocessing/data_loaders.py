import PyPDF2
import requests
import io
import re
from typing import Optional

from docx import Document
from loguru import logger
from streamlit.runtime.uploaded_file_manager import UploadedFile


def clean_text(text: str) -> str:
    """Clean extracted text by handling common encoding and formatting issues."""
    if not text:
        return ""
        
    # Replace common problematic characters
    text = text.replace('\x00', '')  # Remove null bytes
    
    # Fix common encoding artifacts
    text = re.sub(r'[^\x00-\x7F\u0080-\u00FF\u0100-\u017F\u0180-\u024F\u1E00-\u1EFF]', ' ', text)
    
    # Remove multiple spaces and newlines
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text.strip()


def load_pdf(pdf_file: str | UploadedFile) -> Optional[str]:
    """
    Load and extract text from a PDF file with improved encoding handling.
    
    Args:
        pdf_file: Path to PDF file or uploaded file object
        
    Returns:
        Extracted and cleaned text from the PDF, or None if extraction fails
    """
    try:
        if isinstance(pdf_file, UploadedFile):
            pdf_bytes = io.BytesIO(pdf_file.getvalue())
            pdf_reader = PyPDF2.PdfReader(pdf_bytes)
        else:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
        text_parts = []
        for page in pdf_reader.pages:
            try:
                # Extract text from each page
                page_text = page.extract_text()
                if page_text:
                    # Clean the extracted text
                    cleaned_text = clean_text(page_text)
                    if cleaned_text:
                        text_parts.append(cleaned_text)
            except Exception as page_error:
                logger.warning(f"Error extracting text from page: {page_error}")
                continue
                
        if not text_parts:
            logger.warning("No text could be extracted from the PDF")
            return None
            
        return "\n\n".join(text_parts)
        
    except Exception as e:
        logger.exception(f"Error loading PDF: {e}")
        return None


def load_txt(txt_file: str | UploadedFile) -> str | None:
    try:
        if isinstance(txt_file, UploadedFile):
            return txt_file.getvalue().decode("utf-8")
        else:
            with open(txt_file, "r") as file:
                return file.read()
    except Exception as e:
        logger.exception(e)
        return None


def load_docx(docx_file: str | UploadedFile) -> str | None:
    try:
        if isinstance(docx_file, UploadedFile):
            # Convert Streamlit's UploadedFile to bytes IO
            docx_bytes = io.BytesIO(docx_file.getvalue())
            docx_reader = Document(docx_bytes)
        else:
            docx_reader = Document(docx_file)
        return "\n".join(paragraph.text for paragraph in docx_reader.paragraphs)
    except Exception as e:
        logger.exception(e)
        return None


def load_url(url: str) -> str | None:
    try:
        response = requests.get(url)
        response.raise_for_status()
        return response.text
    except Exception as e:
        logger.exception(e)
        return None
